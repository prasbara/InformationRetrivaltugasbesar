"""
Module: loader.py
Purpose: Scans a directory and loads documents of various formats (PDF, DOCX, TXT, MD) into structured LangChain Document objects.
Inputs: File path or directory path (string).
Outputs: Tuple of loaded Document list and load status metadata list.
Workflow: Scans folders, handles extension-specific parsing, extracts text page-by-page (for PDF) or paragraphs (for DOCX), structures metadata, and logs success/failures.
Dependencies: os, typing, langchain_core.documents, pypdf, docx, src.logger.
Complexity: Time: O(P) where P is total number of pages/files processed; Space: O(T) where T is total characters loaded into RAM.
"""
import os
from typing import List, Dict, Any, Tuple
from langchain_core.documents import Document
from pypdf import PdfReader
import docx

from src.logger import log_system

def load_single_document(file_path: str) -> List[Document]:
    """Loads a single document based on its extension and returns a list of Documents."""
    filename = os.path.basename(file_path)
    ext = os.path.splitext(filename)[1].lower()
    documents = []

    try:
        if ext == ".pdf":
            reader = PdfReader(file_path)
            total_pages = len(reader.pages)
            if total_pages == 0:
                log_system(f"PDF file is empty: {filename}", "warning")
                return []
                
            for page_idx, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                # Create a Document for each page to maintain page numbers
                metadata = {
                    "source": filename,
                    "file_path": file_path,
                    "page": page_idx + 1,
                    "total_pages": total_pages,
                    "file_type": "pdf"
                }
                documents.append(Document(page_content=text, metadata=metadata))
                
        elif ext == ".docx":
            doc = docx.Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs])
            metadata = {
                "source": filename,
                "file_path": file_path,
                "page": 1,
                "total_pages": 1,
                "file_type": "docx"
            }
            documents.append(Document(page_content=text, metadata=metadata))
            
        elif ext in [".txt", ".md"]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            metadata = {
                "source": filename,
                "file_path": file_path,
                "page": 1,
                "total_pages": 1,
                "file_type": ext[1:]
            }
            documents.append(Document(page_content=text, metadata=metadata))
        else:
            log_system(f"Unsupported file format: {ext} for file {filename}", "warning")
            
    except Exception as e:
        log_system(f"Error loading file {filename}: {str(e)}", "error")
        # Return empty list; error handling is propagated through empty documents count
        
    return documents

def load_all_documents(directory_path: str) -> Tuple[List[Document], List[Dict[str, Any]]]:
    """Scans directory_path for documents, loads them and returns documents list and status details."""
    all_documents = []
    load_status = []
    
    if not os.path.exists(directory_path):
        os.makedirs(directory_path, exist_ok=True)
        log_system(f"Created documents directory: {directory_path}", "info")
        return [], []
        
    for filename in os.listdir(directory_path):
        file_path = os.path.join(directory_path, filename)
        if os.path.isdir(file_path):
            continue
            
        ext = os.path.splitext(filename)[1].lower()
        if ext not in [".pdf", ".docx", ".txt", ".md"]:
            continue
            
        file_size = os.path.getsize(file_path)
        docs = load_single_document(file_path)
        
        if docs:
            char_count = sum(len(d.page_content) for d in docs)
            pages = max(d.metadata.get("page", 1) for d in docs)
            all_documents.extend(docs)
            load_status.append({
                "filename": filename,
                "status": "Success",
                "size_bytes": file_size,
                "pages": pages,
                "char_count": char_count,
                "error": None
            })
            log_system(f"Successfully loaded {filename} ({pages} pages, {char_count} chars)", "info")
        else:
            load_status.append({
                "filename": filename,
                "status": "Failed",
                "size_bytes": file_size,
                "pages": 0,
                "char_count": 0,
                "error": "File corrupted or unreadable text. Convert to text if PDF is scanned."
            })
            log_system(f"Failed to load document: {filename}", "error")
            
    return all_documents, load_status
